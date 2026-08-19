from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path("FSD-original.docx")
OUTPUT = Path("FSD-Automated-Payment-System-v1.2.docx")


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(9)


def add_row(table, values):
    cells = table.add_row().cells
    for cell, value in zip(cells, values):
        set_cell(cell, str(value))
    return cells


def replace_paragraph(doc, startswith, text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            p.text = text
            for r in p.runs:
                r.font.name = "Arial"
            return p
    raise ValueError(f"Paragraph not found: {startswith}")


def add_before(anchor, text, style=None, bold_prefix=None):
    p = anchor.insert_paragraph_before(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text.split(":", 1)
        p.add_run(first + ":").bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


doc = Document(SOURCE)

# Google Docs exports can contain decimal DXA widths that Word accepts but
# python-docx cannot parse while cloning rows. Normalize them to integer DXA.
for table in doc.tables:
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is not None:
        for grid_col in tbl_grid.gridCol_lst:
            raw = grid_col.get(qn("w:w"))
            if raw and "." in raw:
                grid_col.set(qn("w:w"), str(round(float(raw))))

# Document control/version history.
version_table = doc.tables[1]
version_row = version_table.rows[2].cells
for cell, value in zip(
    version_row,
    ["1.1", "08/19/2026", "Paolo Ylag", "Updated to reflect implemented workflows and approved enhancements", "For stakeholder review"],
):
    set_cell(cell, value)

backend_version_row = version_table.rows[3].cells
for cell, value in zip(
    backend_version_row,
    ["1.2", "08/19/2026", "Paolo Ylag", "Expanded scope for backend-enabled prototype and future implementation roadmap", "For stakeholder review"],
):
    set_cell(cell, value)

# Introduction and scope.
replace_paragraph(
    doc,
    "The system covers reimbursements",
    "The system covers Reimbursement, Liquidation, Cash Advance, P.O. Payment, and General Payment requests. It is intended to improve visibility, reduce follow-ups, standardize approval controls, support role-specific work queues, and maintain a searchable audit trail for Finance.",
)
replace_paragraph(
    doc,
    "The scope of this project is to create",
    "The scope of this project is to create a backend-enabled automated payment request module for Life OS. In addition to the responsive user interface, the prototype includes a FastAPI service and PostgreSQL data model for persistent requests, expense lines, allocations, documents, workflow assignments, approvals, Finance validation, vouchers, payments, notifications, reports, configuration, and immutable audit history. The module allows requestors to select a payment type, save and resume drafts, provide type-specific details, upload supporting documents, and submit completed requests. It routes requests through department, Finance, executive, and Board approvals as applicable; supports tax and accounting review, voucher creation, payment authorization, stakeholder notifications, tracking, reporting, and archival.",
)
replace_paragraph(
    doc,
    "In scope:",
    "In scope: backend APIs and persistence for Reimbursement, Liquidation, Cash Advance, P.O. Payment, and General Payment; authentication context and role-based access; draft save, auto-save, continuation, and deletion; request number generation upon submission; type-specific server-side validation; multiple line items and allocations; currency-aware totals; secure document upload and review; versioned approval routing; controlled request correction and unlocking; Finance tax and accounting validation; voucher generation and lifecycle; payment attempts and signatory authorization; notification queue and delivery history; role-scoped dashboards; tracker, aging, CSV/print reporting, unclaimed-check reporting, archive search, immutable auditing, configurable policies, migrations, seed data, API documentation, automated tests, logging, health checks, and Docker-based local operation.",
)
replace_paragraph(
    doc,
    "Out of scope for the current prototype:",
    "Out of scope for the current backend-enabled prototype: Petty Cash and Credit Card Payment modules; production Life OS single sign-on; live ERP, procurement, banking, and DigiBanker transaction execution; production electronic signatures; production-grade external document storage; historical data migration; and final retention/disaster-recovery infrastructure unless approved for a later phase. The prototype will expose integration-ready interfaces and use local or development substitutes for these external services.",
)

# Actor descriptions.
roles = doc.tables[5]
role_updates = {
    "Requestor": ("Create and save drafts, submit requests, upload documents, correct returned requests, track own requests", "May edit before Document Validation; later changes require an authorized, audited unlock."),
    "Department Head": ("Review request details and documents; approve, return for correction, or decline", "First approval after submission; decisions require an audit entry and reviewer note where applicable."),
    "Finance Associate": ("Validate documents and line items, classify VAT/EWT, prepare accounting entries and vouchers, update payment status", "Validation completion date is system-generated; Check Number is optional and captured during voucher creation."),
    "Finance Manager": ("Review budget status, tax/accounting treatment, approve within authority, and route by threshold", "Confirms whether the transaction is budgeted or unbudgeted before executive routing."),
    "Vendor/Payee": ("Receives processing, pick-up, release, and completion notifications", "No portal access is required in the prototype; notifications include transaction references."),
}
for row in roles.rows[1:]:
    key = row.cells[0].text.strip()
    if key in role_updates:
        set_cell(row.cells[2], role_updates[key][0])
        set_cell(row.cells[3], role_updates[key][1])

# Use case detail.
use_case = doc.tables[6]
uc_values = {
    "Goal": "Complete a controlled payment request from draft through final release and archival.",
    "Trigger": "Requestor selects a payment type or resumes an existing draft.",
    "Primary Flow": "1. Select request type. 2. Enter request and line-item details. 3. Upload required documents. 4. Save draft or submit. 5. Department Head reviews. 6. Finance validates documents, tax, and accounting. 7. Finance Manager confirms budget routing. 8. Required executive/Board approvals are completed. 9. Finance creates the voucher and payment instruction. 10. Authorized signatories approve. 11. Finance records release/completion and the system updates notifications, tracker, and archive.",
    "Alternate Flow": "The assigned reviewer may return the request for correction with comments or decline it with a reason. Authorized Finance personnel may unlock a request after validation; the system records the actor, reason, date/time, and version. Missing documents, unbalanced accounting entries, or incomplete approvals block progression.",
    "Special Requirements": "Role-scoped access; Philippine Time audit timestamps; immutable approval history; printable voucher; currency-aware amounts; responsive and accessible interface.",
    "Open Questions": "Production authentication, document retention period, email provider, ERP/bank integration, and master-data ownership require confirmation before implementation.",
}
for row in use_case.rows:
    label = row.cells[0].text.strip()
    if label in uc_values:
        set_cell(row.cells[1], uc_values[label])

# Functional requirements: preserve existing rows and add implemented behavior.
fr = doc.tables[7]
new_requirements = [
    ("APS-FR-016", "Draft lifecycle", "System supports save, auto-save, resume, and delete; permanent request reference is assigned only upon submission."),
    ("APS-FR-017", "Validation preview", "Preview updates from completed fields, line items, acknowledgements, and uploaded documents."),
    ("APS-FR-018", "Currency handling", "System supports PHP, USD, EUR, and a custom currency while preserving currency context in totals and displays."),
    ("APS-FR-019", "Controlled request editing", "Requestor may edit before Document Validation; later edits require an authorized unlock with reason and audit record."),
    ("APS-FR-020", "Tax computation", "Finance validation follows Total Sales -> 12% VAT -> Net of VAT -> EWT -> Total Amount; transactions at or below PHP 3,000 default to No EWT."),
    ("APS-FR-021", "Line-item validation", "Validation includes Vendor/Merchant, expense account, department/cost center, amount, attachment, and review status per line."),
    ("APS-FR-022", "Validation completion", "The system records the Document Validation Completion Date automatically when all validation controls pass."),
    ("APS-FR-023", "Voucher payment method", "Voucher supports Check, Bank Transfer (DigiBanker), and Cash; Check Number is optional and available only after required approvals."),
    ("APS-FR-024", "Transaction number", "Finance manually records the transaction number, which must correspond to the Finance Team Tracker reference."),
    ("APS-FR-025", "Digital approval certification", "Voucher displays the complete system-verified approval trail, decisions, actors, dates/times, and approval IDs."),
    ("APS-FR-026", "Dashboard and live requests", "Dashboard metrics are clickable; live request selection displays transaction details alongside the list without a separate entry action."),
    ("APS-FR-027", "Reporting and exports", "Finance can filter by department and export an organized Excel-compatible CSV or print/PDF transaction report."),
    ("APS-FR-028", "Aging and overdue status", "Live request and tracker views show aging days and visually identify overdue transactions."),
    ("APS-FR-029", "Payment release notifications", "System records and notifies the requestor/vendor when a payment is ready, available for pick-up, released, returned, declined, or completed."),
    ("APS-FR-030", "Responsive role views", "Role-scoped navigation and queues support mobile drawer, compact desktop, full desktop, light theme, and dark theme layouts."),
    ("APS-FR-031", "Cash Advance dates", "Cash Advance Date equals Event Date and Date to Liquidate is calculated as 15 days after the event end date."),
    ("APS-FR-032", "Liquidation controls", "Liquidation provides a Proof of Return upload for excess cash and automatically carries the related Cash Advance request date."),
    ("APS-FR-033", "P.O. reference", "P.O. Payment selects an approved P.O. reference and displays requestor, supplier, department, amount, and conditional BIR 2303 requirements."),
    ("APS-FR-034", "Tracker navigation", "Tracker keeps key identifiers/status columns visible while users review wide transaction timelines."),
    ("APS-BE-001", "Authentication and RBAC", "API resolves the acting user and enforces role, department, ownership, workflow-assignment, and segregation-of-duties permissions on every protected operation."),
    ("APS-BE-002", "Persistent request API", "FastAPI endpoints persist drafts, submitted requests, request versions, type-specific details, line items, cost-center allocations, and status history in PostgreSQL."),
    ("APS-BE-003", "Server-side business validation", "The API independently validates required fields/documents, totals, allocations, dates, currencies, request-type rules, and permitted state transitions."),
    ("APS-BE-004", "Workflow engine", "Submission creates a versioned approval route; the API controls assignments, sequencing, approve/return/decline/delegate/reroute actions, and current-owner queues."),
    ("APS-BE-005", "Document service", "API supports authorized upload, download, replacement, versioning, metadata, file restrictions, request/line attachment, and Finance review while storing file content outside PostgreSQL."),
    ("APS-BE-006", "Finance validation service", "API persists document decisions, VAT/EWT values, saved tax snapshots, receipt-copy status, accounting entries, balancing controls, reviewer notes, and completion timestamp."),
    ("APS-BE-007", "Voucher lifecycle", "API creates a voucher only after final approval and supports numbering, approved snapshots, posting, printing data, voiding, replacement, and digital approval certification."),
    ("APS-BE-008", "Payment lifecycle", "API records payment attempts, methods, transaction/check references, preparation, signatory authorization, pick-up availability, release, clearing, completion, failure, voiding, and replacement."),
    ("APS-BE-009", "Notification service", "Business events create idempotent notification jobs with recipient, template, payload, queue/sent/retry/failure state, attempt history, and development delivery adapter."),
    ("APS-BE-010", "Dashboard and reporting API", "API provides role-scoped metrics, queues, filters, pagination, sorting, aging/overdue calculations, tracker data, archive search, CSV export, and print-report datasets."),
    ("APS-BE-011", "Immutable audit service", "Material actions store actor/system identity, action, entity, reason, request version, before/after values where applicable, correlation identifier, and Philippine Time timestamp."),
    ("APS-BE-012", "Configuration and policy versioning", "Approval thresholds, document rules, tax codes, cost centers, payment methods, numbering, notification templates, and aging rules are seeded/configurable with effective dates and history."),
    ("APS-BE-013", "Idempotency and transactions", "Submission, approvals, voucher/payment creation, integration calls, and notifications use database transactions and idempotency keys to prevent duplicate financial actions."),
    ("APS-BE-014", "Operational API controls", "Backend includes standardized errors, structured logs, correlation IDs, health/readiness endpoints, secure secrets handling, query limits, and safe masking of financial data."),
    ("APS-BE-015", "Database delivery", "Alembic migrations create the complete schema, constraints, indexes, seed configuration, and demonstration users/data; destructive history deletion is unavailable through normal APIs."),
    ("APS-BE-016", "Automated verification", "Tests cover request types, permissions, validation failures, threshold boundaries, workflow exceptions, tax/accounting controls, idempotency, reporting, and end-to-end payment completion."),
    ("APS-BE-017", "Integration adapters", "Interfaces isolate Life OS identity, P.O./vendor master data, email, file storage, ERP, and banking so development adapters can later be replaced without rewriting workflow logic."),
    ("APS-BE-018", "Backup and recovery readiness", "Prototype documents database backup/restore, migration rollback, seed/reset, and recovery procedures; production recovery objectives remain subject to infrastructure approval."),
]
for row in new_requirements:
    add_row(fr, row)

# Field-level additions.
fields = doc.tables[8]
field_rows = [
    (12, "Draft Status", "System status", "Yes", "No", "Text", "Draft; Submitted; Returned; Declined; Completed", "Draft", "Draft", "System generated"),
    (13, "Currency", "Dropdown/custom entry", "Yes", "Yes", "Text", "PHP; USD; EUR; Custom", "PHP", "USD", "User entry"),
    (14, "Vendor/Merchant", "Text per line item", "Conditional", "Yes", "Text", "None", "None", "Metro Repairs", "User entry"),
    (15, "VAT Classification", "Dropdown", "Conditional", "Yes", "Text", "VAT; Non-VAT; Zero-rated", "None", "VAT", "Finance Associate"),
    (16, "EWT Classification", "Dropdown", "Conditional", "Yes", "Text", "No EWT; configured rates", "No EWT when amount <= PHP 3,000", "2%", "Finance Associate/system rule"),
    (17, "Validation Completion Date", "System date/time", "Conditional", "No", "DateTime", "Philippine Time", "None", "08/19/2026 14:30", "System generated"),
    (18, "Payment Method", "Dropdown", "Yes before payment", "Yes", "Text", "Check; Bank Transfer (DigiBanker); Cash", "None", "Bank Transfer (DigiBanker)", "Finance Associate"),
    (19, "Check Number", "Text input", "No", "Yes", "Text", "None", "Blank", "CHK-001245", "Finance Associate"),
    (20, "Transaction Number", "Text input", "Yes before release", "Yes", "Text", "Finance tracker reference", "None", "FTT-2026-00891", "Finance Associate"),
    (21, "Proof of Return", "File upload", "Conditional", "Yes", "File", "PDF/Image/Office files", "None", "excess-cash-return.pdf", "Requestor"),
]
for row in field_rows:
    add_row(fields, row)

# Business rules.
rules = doc.tables[9]
rule_rows = [
    ("Draft submission", "Permanent reference is generated only when the completed request is submitted.", "Complete all required fields and documents before submission.", "Draft and validation state", "Auto-save does not submit the request."),
    ("Tax calculation", "Compute Total Sales, 12% VAT, Net of VAT, EWT, then Total Amount.", "Review the VAT/EWT classification and computation.", "Amount and tax classification", "No EWT applies automatically at or below PHP 3,000."),
    ("Accounting entries", "Total debits must equal total credits before validation can be completed.", "Accounting entries must be balanced.", "Line items and tax calculation", "Completion date is recorded automatically after validation passes."),
    ("Check Number", "May be entered only after required approvals and only when Payment Method is Check.", "Complete required approvals before entering a check number.", "Approval trail and payment method", "Optional for DigiBanker and Cash transactions."),
    ("Cash Advance liquidation", "Date to Liquidate equals event end date plus 15 calendar days.", "Enter a valid event end date.", "Event dates", "Cash Advance Date equals Event Date."),
    ("Returned request", "Requestor may correct and resubmit; system retains reviewer comments and version history.", "Address the reviewer comments before resubmitting.", "Workflow and audit trail", "Later edits may require an authorized unlock."),
]
for row in rule_rows:
    add_row(rules, row)

# Buttons/actions.
actions = doc.tables[10]
action_rows = [
    ("Save Draft", "Saves current request state", "Auto-save may also run", "Before submission", "Enabled when request is editable", "My Drafts", "None", "Request ownership"),
    ("Delete Draft", "Deletes selected draft after confirmation", "None", "Drafts only", "Enabled for draft owner", "My Drafts", "Confirmation required", "Draft status"),
    ("Return for Correction", "Returns request with reviewer comments", "Sends notification", "Assigned reviewers", "Enabled for current reviewer", "Requestor correction view", "Reviewer comment required", "Workflow status"),
    ("Complete Validation", "Records validation completion and routes onward", "Sets completion timestamp", "Finance Associate", "Enabled when documents and entries pass", "Finance Manager queue", "Balanced entries and required documents", "Validation state"),
    ("Mark Available for Pick-up", "Records availability and Finance actor", "Sends requestor/vendor notifications", "Finance processing", "Enabled after authorization", "Payment tracker", "Payment reference required", "Authorization status"),
    ("Export CSV", "Exports filtered department transactions", "None", "Finance reports", "Enabled when results exist", "File download", "None", "Report filters"),
]
for row in action_rows:
    add_row(actions, row)

# Configuration, non-functional, and reporting sections.
replace_paragraph(
    doc,
    "The application requires configuration",
    "The application requires configuration for payment request types; draft retention; common and type-specific fields; required and conditional documents; currencies; VAT/EWT rules; approval thresholds; workflow owners; role permissions; department/cost-center values; email recipients and templates; voucher numbering and payment methods; status and aging rules; tracker columns; audit retention; and archive/search permissions.",
)

# Backend implementation scope, inserted after detailed UI specifications and
# before system configuration requirements.
system_config_anchor = next(p for p in doc.paragraphs if p.text.strip() == "System Configurations" and p._p.getparent() is doc.element.body)
backend_heading = add_before(system_config_anchor, "3.1.6 Backend Implementation Scope")
backend_heading.runs[0].bold = True
backend_heading.runs[0].font.name = "Crimson Pro"
backend_heading.runs[0].font.size = Pt(14)
backend_scope_paragraphs = [
    ("Architecture", "FastAPI exposes versioned REST endpoints backed by PostgreSQL and Alembic migrations. Business services own validation and workflow rules; integration adapters isolate identity, files, email, purchasing, ERP, and banking dependencies."),
    ("Identity and authorization", "The prototype supports demonstration identities or integration-ready Life OS identity claims. The API enforces role, ownership, department, assignment, and segregation-of-duties rules rather than relying on hidden front-end controls."),
    ("Request domain", "Persistent entities include users, roles, departments, cost centers, vendors/payees, requests, type-specific request details, expense lines, allocations, documents, request versions, workflow assignments, and status history."),
    ("Workflow", "Submission snapshots the applicable policy and creates the approval route. State transitions are server-controlled and support approval, return, decline, delegation, authorized rerouting, correction, resubmission, cancellation, and audited unlocking."),
    ("Finance and payment domain", "The backend stores document decisions, tax snapshots, accounting entries, vouchers, digital approval certifications, payment attempts, payment methods, signatory decisions, check/bank references, pick-up and release records, and failed or voided attempts."),
    ("Documents", "File content is held in local development storage for the prototype while PostgreSQL stores metadata, ownership, document type, version, review status, hashes, and audit references. Storage is accessed only through authorized API operations."),
    ("Notifications", "Workflow events create durable notification jobs. Delivery status is independent of business status and records queueing, successful delivery, retries, failures, and deduplication keys. A development mailbox/log may substitute for live email."),
    ("Queries and reports", "Backend endpoints provide role-scoped dashboard counts, queues, request detail, workflow timelines, aging/overdue values, tracker results, unclaimed checks, completed payments, archive search, paginated lists, CSV exports, and print-ready datasets."),
    ("Audit and time", "All material actions create immutable audit events. API timestamps use Asia/Manila and include the UTC+08:00 offset; date-only business values remain calendar dates."),
    ("Delivery and quality", "The repository includes Docker-based local operation, environment configuration, seed/demo data, OpenAPI documentation, standardized errors, structured logs, health/readiness endpoints, database backup/restore instructions, and automated unit, integration, permission, workflow, and end-to-end tests."),
]
for label, description in backend_scope_paragraphs:
    add_before(system_config_anchor, f"{label}: {description}", bold_prefix=label)
replace_paragraph(
    doc,
    "- Exportable report format",
    "- Department Transaction Report with department filters, an organized Excel-compatible CSV export, and a print/PDF layout.",
)
report_anchor = next(p for p in doc.paragraphs if p.text.strip().startswith("Data Migration/Conversion Requirements"))
additional_reports = [
    "- Role-scoped dashboard totals and clickable lists for pending, returned, overdue, available-for-pick-up, and completed requests.",
    "- Aging report showing elapsed days and overdue indicators by current workflow owner.",
    "- Payment completion report including payment method, transaction/check reference, release date, and recording Finance personnel.",
]
for text in additional_reports:
    add_before(report_anchor, text)

# Insert missing integration and exception-handling body sections before migration.
integration_heading = add_before(report_anchor, "7. Integration Requirements")
integration_heading.runs[0].bold = True
integration_heading.runs[0].font.name = "Crimson Pro"
integration_heading.runs[0].font.size = Pt(16)
add_before(report_anchor, "The production module is expected to integrate with Life OS authentication and role management, department and cost-center master data, document storage, email delivery, approved P.O. data, and the Finance Team Tracker. ERP posting, banking/DigiBanker connectivity, and vendor-master synchronization remain future integrations until interfaces, ownership, and security requirements are approved.")
exception_heading = add_before(report_anchor, "7.1 Exception Handling and Error Reporting")
exception_heading.runs[0].bold = True
exception_heading.runs[0].font.name = "Crimson Pro"
exception_heading.runs[0].font.size = Pt(14)
add_before(report_anchor, "The system must prevent progression when mandatory fields or documents are missing, accounting entries are unbalanced, an approval is incomplete, a user lacks authority, or a transaction reference is invalid. Errors must use plain-language messages, preserve entered data, and identify the corrective action. Integration failures must be logged with request reference, interface, timestamp in Philippine Time, retry status, and a support-safe error identifier.")

# Open issues: replace blank row with tracked decisions/open production items.
issues = doc.tables[12]
issue_rows = [
    ("APS-OI-001", "Production authentication and role directory", "Project Team", "08/19/2026", "Confirm Life OS identity source and role provisioning.", "", "", "Open"),
    ("APS-OI-002", "Email and notification delivery service", "Finance/DT", "08/19/2026", "Confirm provider, templates, distribution lists, and bounce handling.", "", "", "Open"),
    ("APS-OI-003", "ERP, P.O., bank, and DigiBanker integration scope", "Finance/DT", "08/19/2026", "Define interfaces, ownership, reconciliation, and security controls.", "", "", "Open"),
    ("APS-OI-004", "Retention and archival policy", "Finance", "08/19/2026", "Confirm retention periods for requests, documents, audit events, and reports.", "", "", "Open"),
]
for i, row in enumerate(issue_rows):
    target = issues.rows[1].cells if i == 0 else issues.add_row().cells
    for cell, value in zip(target, row):
        set_cell(cell, value)

# Acceptance criteria.
acceptance = doc.tables[13]
criteria = [
    ("Create and manage drafts", "A requestor can select a request type, save/auto-save, resume, and delete a draft; no permanent request reference exists until submission."),
    ("Submit each supported request type", "Reimbursement, Liquidation, Cash Advance, P.O. Payment, and General Payment enforce their configured fields, documents, and line-item rules."),
    ("Validate documents and accounting", "Finance can review Vendor/Merchant and attachments per line, classify VAT/EWT, balance debits/credits, and complete validation with a system timestamp."),
    ("Route approvals", "The system routes by budget status and amount, including Board approval only for unbudgeted requests above PHP 1,000,000, and retains every decision."),
    ("Return or decline a request", "Reviewer comments, actor, timestamp, request version, notification, and requestor-specific link are retained; corrected requests can be resubmitted."),
    ("Create a payment voucher", "Voucher is available only after final approval and shows payment method, tax/net amounts, transaction/check reference as applicable, accounting entries, and the complete digital approval trail."),
    ("Process and release payment", "Finance can record authorization, processing, pick-up availability, release, and completion with actor/date/time and stakeholder notifications."),
    ("Track and report transactions", "Authorized users can open transaction detail from dashboards/trackers, filter Finance reports, identify aging/overdue items, export CSV, and print/PDF reports."),
    ("Use role-scoped responsive views", "Each actor sees only permitted navigation, requests, actions, and queues across mobile, compact desktop, full desktop, light, and dark layouts."),
    ("Maintain an audit trail", "Draft submission, edits/unlocks, uploads, approvals, validation, voucher changes, payment updates, notifications, and release records are timestamped in Philippine Time and remain searchable."),
]
for i, row in enumerate(criteria):
    target = acceptance.rows[i + 1].cells if i + 1 < len(acceptance.rows) else acceptance.add_row().cells
    set_cell(target[0], row[0])
    set_cell(target[1], row[1])

# Improve table readability and repeat header rows.
for table in doc.tables:
    if table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
        for cell in table.rows[0].cells:
            shade_cell(cell, "8A1538")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xF8, 0xE7)

# Footer version marker without disturbing existing content.
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    if "FSD v1.1" not in p.text:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("FSD v1.2 | Automated Payment System | Updated 19 August 2026").font.size = Pt(8)

# Normalize fractional half-point font sizes produced by Google Docs export.
for element in doc.element.body.xpath(".//w:sz | .//w:szCs"):
    raw = element.get(qn("w:val"))
    if raw and "." in raw:
        element.set(qn("w:val"), str(round(float(raw))))

# Replace legacy green accents with the approved crimson-and-ivory palette.
def is_green(hex_value):
    if not hex_value or len(hex_value) != 6:
        return False
    try:
        r, g, b = (int(hex_value[i:i + 2], 16) for i in (0, 2, 4))
    except ValueError:
        return False
    return g > r * 1.12 and g > b * 1.08 and g - min(r, b) > 18


for element in doc.element.body.iter():
    for attr_name, attr_value in list(element.attrib.items()):
        local = attr_name.split("}")[-1]
        if local in {"color", "fill", "val"} and is_green(attr_value.upper()):
            # Pale background fills become ivory; text, rules, and borders become crimson.
            rgb = tuple(int(attr_value[i:i + 2], 16) for i in (0, 2, 4))
            light = sum(rgb) / 3 > 185
            replacement = "FFF8E7" if local == "fill" and light else "8A1538"
            element.set(attr_name, replacement)

doc.save(OUTPUT)
print(OUTPUT.resolve())
